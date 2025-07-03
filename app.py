import os
import shutil
import zipfile
import pandas as pd
from flask import Flask, request, send_file, render_template, redirect, url_for, flash, jsonify
from lxml import etree
from lxml.etree import QName
import re
import subprocess
import tempfile
from docx import Document
from werkzeug.utils import secure_filename
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uuid
import datetime
import requests
import numpy as np
# 并发处理模块
import concurrent.futures
import threading
from queue import Queue
import time
import random
from collections import Counter


app = Flask(__name__)
app.secret_key = 'replace-with-your-own-secret'

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
def W(tag): return f"{{{W_NS}}}{tag}"

DEEPSEEK_API_KEY = 'sk-44707dc99db4416c9d21260fb0d9f6bc'
DEEPSEEK_API_URL = 'https://api.deepseek.com/v1/chat/completions'

# 并发处理参数置

CONCURRENT_CONFIG = {
    'enabled': True,           # 是否启用并发处理
    'batch_size': 3,           # 每批处理的段落数量（从3增加到5）
    'max_workers': 3,          # 最大并发线程数（从3增加到5）
    'threshold': 0.5           # 相似度阈值
}

# 线程锁，用于保护XML操作
xml_lock = threading.Lock()

# 动态调整并发参数
def adjust_concurrent_config(total_paragraphs, total_comments):
    """
    根据文档大小动态调整并发参数
    """
    print(f"\n=== 动态调整并发参数 ===")
    print(f"文档段落数: {total_paragraphs}")
    print(f"批注数量: {total_comments}")
    
    # 根据文档大小调整批次大小
    if total_paragraphs <= 10:
        CONCURRENT_CONFIG['batch_size'] = 3
        CONCURRENT_CONFIG['max_workers'] = 4
        print(f"  小文档：批次大小调整为 {CONCURRENT_CONFIG['batch_size']}，并发数调整为 {CONCURRENT_CONFIG['max_workers']}")
    elif total_paragraphs <= 30:
        CONCURRENT_CONFIG['batch_size'] = 5
        CONCURRENT_CONFIG['max_workers'] = 6
        print(f"  中等文档：批次大小调整为 {CONCURRENT_CONFIG['batch_size']}，并发数调整为 {CONCURRENT_CONFIG['max_workers']}")
    elif total_paragraphs <= 100:
        CONCURRENT_CONFIG['batch_size'] = 8
        CONCURRENT_CONFIG['max_workers'] = 8
        print(f"  大文档：批次大小调整为 {CONCURRENT_CONFIG['batch_size']}，并发数调整为 {CONCURRENT_CONFIG['max_workers']}")
    else:
        CONCURRENT_CONFIG['batch_size'] = 10
        CONCURRENT_CONFIG['max_workers'] = 10
        print(f"  超大文档：批次大小调整为 {CONCURRENT_CONFIG['batch_size']}，并发数调整为 {CONCURRENT_CONFIG['max_workers']}")
    
    # 根据批注数量调整阈值
    if total_comments <= 5:
        CONCURRENT_CONFIG['threshold'] = 0.3
        print(f"  少量批注：阈值调整为 {CONCURRENT_CONFIG['threshold']}")
    elif total_comments <= 20:  
        CONCURRENT_CONFIG['threshold'] = 0.5
        print(f"  中等批注：阈值调整为 {CONCURRENT_CONFIG['threshold']}")
    else:
        CONCURRENT_CONFIG['threshold'] = 0.6
        print(f"  大量批注：阈值调整为 {CONCURRENT_CONFIG['threshold']}")
    
    print(f"✓ 并发参数调整完成")

def extract_question_numbers_and_classify(text):
    """
    从文档文本中提取题号并进行分类
    返回: (题号列表, 分类结果)
    """
    print(f"\n=== 开始提取题号并分类 ===")
    
    # 匹配常见的题号格式
    question_patterns = [
        r'(\d+)[\.、]\s*',  # 1. 或 1、
        r'（(\d+)）\s*',    # （1）
        r'\((\d+)\)\s*',    # (1)
        r'第(\d+)题\s*',    # 第1题
        r'(\d+)\.\s*',      # 1.
    ]
    
    question_numbers = []
    lines = text.split('\n')
    
    for line_num, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        for pattern in question_patterns:
            match = re.match(pattern, line)
            if match:
                question_num = int(match.group(1))
                question_numbers.append({
                    'number': question_num,
                    'line': line_num,
                    'text': line,
                    'full_text': line
                })
                break
    
    print(f"找到题号: {[q['number'] for q in question_numbers]}")
    
    # 使用AI进行分类
    if question_numbers:
        return classify_questions_with_ai(question_numbers, text)
    else:
        return [], []

def classify_questions_with_ai(question_numbers, full_text):
    """
    使用AI对题号进行分类，生成分类标题
    """
    print(f"\n=== 使用AI进行题号分类 ===")
    
    headers = {
        'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
        'Content-Type': 'application/json',
    }
    
    # 构建题号信息
    question_info = []
    for q in question_numbers:
        # 获取题号后的内容（最多200字符）
        start_pos = len(q['text'])
        end_pos = min(start_pos + 200, len(full_text))
        content = full_text[start_pos:end_pos].strip()
        question_info.append(f"题号{q['number']}: {content}")
    
    system_prompt = """你是一个专业的文档分析专家。你的任务是对文档中的题号进行划分，并为每个分类生成合适的标题。

请分析题号的内容，将它们按照主题、类型或内容相似性进行分组，并为每个分组生成一个简洁的标题。

请严格按照以下JSON格式返回结果：
{
    "categories": [
        {
            "title": "分类标题",
            "question_numbers": [题号列表],
            "description": "该分类的简要描述"
        }
    ]
}"""

    user_prompt = f"""文档中的题号信息：
{chr(10).join(question_info)}

请对这些题号进行分类，并生成合适的分类标题。"""

    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 800,
        "response_format": {"type": "json_object"}
    }
    
    # 增加重试机制
    max_retries = 3
    retry_delay = 2  # 秒
    
    for attempt in range(max_retries):
        try:
            # 使用更长的超时时间和重试设置
            session = requests.Session()
            session.mount('https://', requests.adapters.HTTPAdapter(
                max_retries=requests.adapters.Retry(
                    total=3,
                    backoff_factor=0.5,
                    status_forcelist=[500, 502, 503, 504]
                )
            ))
            
            response = session.post(
                DEEPSEEK_API_URL, 
                headers=headers, 
                json=data, 
                timeout=(10, 60),  # (连接超时, 读取超时)
                verify=True
            )
            
            if response.status_code != 200:
                print(f'  ❌ DeepSeek API error: {response.status_code} - {response.text}')
                if attempt < max_retries - 1:
                    print(f'  ⏳ 重试中... ({attempt + 1}/{max_retries})')
                    time.sleep(retry_delay)
                    continue
                return None
            
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            
            import json
            classification = json.loads(content)
            
            categories = classification.get('categories', [])
            print(f"AI分类结果: {len(categories)} 个分类")
            for cat in categories:
                print(f"  - {cat['title']}: {cat['question_numbers']}")
            
            return question_numbers, categories
            
        except requests.exceptions.ConnectionError as e:
            print(f'  ❌ 连接错误 (尝试 {attempt + 1}/{max_retries}): {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2  # 指数退避
                continue
            else:
                print(f'  ❌ 连接失败，已达到最大重试次数')
                return None
                
        except requests.exceptions.Timeout as e:
            print(f'  ❌ 请求超时 (尝试 {attempt + 1}/{max_retries}): {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                print(f'  ❌ 请求超时，已达到最大重试次数')
                return None
                
        except Exception as e:
            print(f'  ❌ AI分类失败: {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                print(f'  ❌ API调用失败，已达到最大重试次数')
                return None

def add_title_comments_to_docx(docx_path, categories, output_path):
    """
    在docx文档中为分类标题添加批注
    """
    print(f"\n=== 添加分类标题批注 ===")
    
    temp_dir = 'temp_docx_' + str(uuid.uuid4())
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
    except Exception as e:
        raise Exception(f"文件解压失败: {str(e)}")

    doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
    comments_xml_path = os.path.join(temp_dir, 'word', 'comments.xml')
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parser = etree.XMLParser(remove_blank_text=True, recover=True)
    doc_tree = etree.parse(doc_xml_path, parser)
    doc_root = doc_tree.getroot()
    paragraphs = doc_root.findall('.//w:p', namespaces={'w': w_ns})

    # 创建或加载comments.xml
    if os.path.exists(comments_xml_path):
        comments_tree = etree.parse(comments_xml_path, parser)
        comments_root = comments_tree.getroot()
    else:
        MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
        W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
        WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
        nsmap = {'w': w_ns, 'mc': MC_NS, 'w14': W14_NS, 'w15': W15_NS, 'wp14': WP14_NS}
        comments_root = etree.Element(f'{{{w_ns}}}comments', nsmap=nsmap)
        comments_root.set(f"{{{MC_NS}}}Ignorable", "w14 w15 wp14")
        comments_tree = etree.ElementTree(comments_root)

    current_id = 0
    
    # 为每个分类添加标题批注
    for category in categories:
        title = category['title']
        question_nums = category['question_numbers']
        description = category.get('description', '')
        
        # 查找包含第一个题号的段落
        first_question = question_nums[0] if question_nums else None
        if not first_question:
            continue
            
        # 在文档中查找该题号
        for para_idx, para in enumerate(paragraphs):
            para_text = ''.join(t.text for t in para.findall('.//w:t', namespaces={'w': w_ns}) if t.text)
            para_text = para_text.strip()
            
            # 检查是否包含该题号
            if re.search(fr'\b{first_question}\b', para_text):
                # 在该段落前插入标题批注
                comment_id = str(current_id)
                
                # 创建批注范围
                comment_start_node = etree.Element(W('commentRangeStart'))
                comment_start_node.set(W('id'), comment_id)
                
                # 插入到段落开始
                parent = para.getparent()
                para_index = parent.index(para)
                parent.insert(para_index, comment_start_node)
                
                comment_end_node = etree.Element(W('commentRangeEnd'))
                comment_end_node.set(W('id'), comment_id)
                parent.insert(para_index + 1, comment_end_node)
                
                # 创建批注引用
                comment_ref_run = etree.Element(W('r'))
                comment_ref_node = etree.Element(W('commentReference'))
                comment_ref_node.set(W('id'), comment_id)
                comment_ref_run.append(comment_ref_node)
                parent.insert(para_index + 2, comment_ref_run)
                
                # 创建批注内容
                comment_elem = etree.Element(W('comment'))
                comment_elem.set(W('id'), comment_id)
                comment_elem.set(W('author'), '分类系统')
                comment_elem.set(W('date'), datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))

                # 标题
                p1 = etree.Element(W('p'))
                r1 = etree.Element(W('r'))
                t1 = etree.Element(W('t'))
                t1.text = clean_comment_text(f"分类标题：{title}")
                r1.append(t1)
                p1.append(r1)
                comment_elem.append(p1)

                # 包含的题号
                p2 = etree.Element(W('p'))
                r2 = etree.Element(W('r'))
                t2 = etree.Element(W('t'))
                t2.text = clean_comment_text(f"包含题号：{', '.join(map(str, question_nums))}")
                r2.append(t2)
                p2.append(r2)
                comment_elem.append(p2)

                # 描述
                if description:
                    p3 = etree.Element(W('p'))
                    r3 = etree.Element(W('r'))
                    t3 = etree.Element(W('t'))
                    t3.text = clean_comment_text(f"描述：{description}")
                    r3.append(t3)
                    p3.append(r3)
                    comment_elem.append(p3)

                comments_root.append(comment_elem)
                print(f"  ✓ 为分类 '{title}' 添加了标题批注")
                current_id += 1
                break
    
    # 保存XML
    doc_tree.write(doc_xml_path, xml_declaration=True, encoding='utf-8', standalone='yes')
    comments_tree.write(comments_xml_path, xml_declaration=True, encoding='utf-8', standalone='yes')
    
    # 打包为新的docx文件
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
        for foldername, subfolders, filenames in os.walk(temp_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, temp_dir)
                docx.write(file_path, arcname)

    # 清理临时文件夹
    shutil.rmtree(temp_dir)
    print(f"✓ 分类标题批注已添加到: {output_path}")

# 并发批注匹配处理
def process_paragraph_batch(paragraph_batch, comment_texts, threshold=0.5, max_workers=3):
    """
    并发处理一批段落，同时处理多个段落的批注匹配
    paragraph_batch: 段落批次，每个元素包含 (para_idx, para, para_text, context_text)
    comment_texts: 批注内容列表
    threshold: 相似度阈值
    max_workers: 最大并发数
    """
    print(f"\n=== 开始并发处理 {len(paragraph_batch)} 个段落，最大并发数: {max_workers} ===")
    
    results = []
    
    def process_single_paragraph(para_data):
        para_idx, para, para_text, context_text = para_data
        try:
            # 调用原有的匹配函数
            match = deepseek_context_aware_match(para_text, comment_texts, context_text, threshold)
            if match:
                text_to_find, original_comment, similarity_score, comment_idx, reasoning, context_analysis, error_type, comment_type, is_title, logic_consistent, modify_suggestion = match
                return {
                    'para_idx': para_idx,
                    'para': para,
                    'para_text': para_text,
                    'match': match,
                    'success': True
                }
            else:
                return {
                    'para_idx': para_idx,
                    'para': para,
                    'para_text': para_text,
                    'match': None,
                    'success': False
                }
        except Exception as e:
            # 使用更安全的错误处理
            error_msg = f"处理段落 {para_idx} 时出错: {str(e)}"
            try:
                print(f"  ❌ {error_msg}")
            except:
                pass  # 避免在程序关闭时打印错误
            return {
                'para_idx': para_idx,
                'para': para,
                'para_text': para_text,
                'match': None,
                'success': False,
                'error': str(e)
            }
    
    # 使用线程池执行并发处理，设置线程为非守护线程
    executor = concurrent.futures.ThreadPoolExecutor(
        max_workers=max_workers,
        thread_name_prefix="CommentMatch"
    )
    
    try:
        # 提交所有任务
        future_to_para = {executor.submit(process_single_paragraph, para_data): para_data for para_data in paragraph_batch}
        
        # 收集结果
        for future in concurrent.futures.as_completed(future_to_para):
            para_data = future_to_para[future]
            try:
                result = future.result(timeout=60)  # 添加超时
                results.append(result)
                
                if result['success']:
                    para_idx = result['para_idx']
                    text_to_find, original_comment, similarity_score, comment_idx, reasoning, context_analysis, error_type, comment_type, is_title, logic_consistent, modify_suggestion = result['match']
                    try:
                        print(f"  ✓ 并发处理段落 {para_idx} 匹配成功: '{original_comment[:50]}...' (相似度: {similarity_score:.3f})")
                    except:
                        pass  # 避免在程序关闭时打印错误
                else:
                    para_idx = result['para_idx']
                    try:
                        print(f"  ⚠ 并发处理段落 {para_idx} 未匹配到批注")
                    except:
                        pass
                    
            except concurrent.futures.TimeoutError:
                para_idx = para_data[0]
                try:
                    print(f"  ❌ 并发处理段落 {para_idx} 超时")
                except:
                    pass
                results.append({
                    'para_idx': para_idx,
                    'para': para_data[1],
                    'para_text': para_data[2],
                    'match': None,
                    'success': False,
                    'error': '处理超时'
                })
            except Exception as e:
                para_idx = para_data[0]
                try:
                    print(f"  ❌ 并发处理段落 {para_idx} 时发生异常: {str(e)}")
                except:
                    pass
                results.append({
                    'para_idx': para_idx,
                    'para': para_data[1],
                    'para_text': para_data[2],
                    'match': None,
                    'success': False,
                    'error': str(e)
                })
    
    finally:
        # 确保线程池正确关闭
        executor.shutdown(wait=True)
    
    # 按原始段落索引排序结果
    results.sort(key=lambda x: x['para_idx'])
    
    try:
        print(f"✓ 并发处理完成，成功处理 {len([r for r in results if r['success']])} 个段落")
    except:
        pass  # 避免在程序关闭时打印错误
    
    return results

def create_paragraph_batches(paragraphs, para_to_block_idx, comment_texts, batch_size=3):
    """
    将段落分批，为并发处理做准备
    """
    print(f"\n=== 创建段落批次，批次大小: {batch_size} ===")
    
    batches = []
    current_batch = []
    
    for para_idx, para in enumerate(paragraphs):
        para_text = ''.join(t.text for t in para.findall('.//w:t', namespaces={'w': W_NS}) if t.text).strip()
        
        # 跳过标题段落（只在内容段落插入批注）
        if para_idx not in para_to_block_idx or not para_text or not comment_texts:
            continue
            
        # 构建同区块下的上下文
        block = para_to_block_idx[para_idx]
        content_indices = block['content_indices']
        idx_in_block = content_indices.index(para_idx)
        context_text = ""
        
        if idx_in_block > 0:
            prev_idx = content_indices[idx_in_block - 1]
            prev_text = ''.join(t.text for t in paragraphs[prev_idx].findall('.//w:t', namespaces={'w': W_NS}) if t.text).strip()
            if prev_text:
                context_text += f"同区块前文：{prev_text}\n"
        if idx_in_block < len(content_indices) - 1:
            next_idx = content_indices[idx_in_block + 1]
            next_text = ''.join(t.text for t in paragraphs[next_idx].findall('.//w:t', namespaces={'w': W_NS}) if t.text).strip()
            if next_text:
                context_text += f"同区块后文：{next_text}"
        
        # 添加到当前批次
        current_batch.append((para_idx, para, para_text, context_text))
        
        # 如果批次满了，保存并开始新批次
        if len(current_batch) >= batch_size:
            batches.append(current_batch)
            current_batch = []
    
    # 添加最后一个不完整的批次
    if current_batch:
        batches.append(current_batch)
    
    print(f"✓ 创建了 {len(batches)} 个批次，总共 {sum(len(batch) for batch in batches)} 个段落")
    return batches



def deepseek_context_aware_match(target_text, doc_texts, context_text="", threshold=0.5):
    """
    上下文感知的DeepSeek批注匹配，优先考虑上下文
    target_text: 当前段落文本
    doc_texts: 批注内容列表
    context_text: 上下文文本（同一标题下的前后段落）
    threshold: 相似度阈值
    """
    print(f"\n=== DeepSeek上下文感知匹配开始 ===")
    print(f"当前段落: '{target_text}'")
    print(f"上下文: '{context_text}'")
    # print(f"批注数量: {len(doc_texts)}")
    
    headers = {
        'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
        'Content-Type': 'application/json',
    }
    
    candidate_comments = [f"{i+1}. {c}" for i, c in enumerate(doc_texts) if c.strip()]
    if not candidate_comments:
        print("  ⚠ 没有找到有效的候选批注")
        return None
    
    system_prompt = """你是一个专业的法务/律师/项目经理，负责为合同文档添加批注。请像真实人类一样写批注，要求如下：
1. 语言自然、简洁、直接，避免模板化。
2. 可以提问、建议、提醒、补充说明。
3. 遇到条款不明确、需补充、需确认的地方，直接提出。
4. 不要机械重复，不要只复述原文，要有针对性和场景感。

重要规则：
1. 精确错误识别：只要发现明显的错误，无论文本多短都要插入批注。
2. 错误类型包括但不限于：日期错误、数字错误、名称错误、格式错误等。
3. 上下文分析：考虑同一标题下的前后段落，帮助判断是否真的存在错误。
4. 语义相关性：段落内容与批注在语义上是否高度相关。
5. 精确性：批注是否针对段落中的特定错误内容，并且批注不要绑定到标题上，批注只与正文相绑定。
6. 仿真性检查：仅对具体错误批注(specific_error)进行逻辑一致性分析。若原文逻辑是正确的与批注内容相符合的，就不要进行批注。例如：批注内容：验收日期定为3日，不可更改。原文内容：验收期为4月3日。这里原文与批注都是3日，逻辑相符合，不要进行批注。全局性批注(global_comment)不进行逻辑一致性分析。
7. target_text_in_paragraph 必须是当前段落中的实际文本，不能是上下文的文本。
8. 全局性批注：如果批注是对整个文档的宏观评价或建议，标记为global_comment类型，找到最适合的段落位置插入，且每个全局批注只应该匹配一次。
9. 标题过滤：如果当前段落是标题（包含数字编号、章节标识等），则不进行批注匹配。
10. 直接使用：将你获得的批注作为prompt，然后根据这个批注找到原文中错误的地方。
11. 如果原文中没有错误不要进行批注。
12. 批注找到插入的地方之后，再加上修改建议.
13. 每条批注内容只能归为一种类型，不能同时是两种类型。
14. 如果一条批注内容针对某一条款、某一段、某一细节（如付款条款、违约条款等），那么就归为具体批注类。
15. 题出批注原因的语言要自然、简洁、直接，避免模板化。
16. 批注内容要符合法律法规，不要出现违法违规的内容。
17. 遇到条款不明确，需要补充、确认的原文，找到合适的批注进行批注，并表明原因，给出修改建议。
18. 批注内容要符合合同条款，不要出现与合同条款不符的内容。
19. 仔细检查文档的内容细节，类似于报价、服务内容、支付方式等联系上下文检查是否有表明，若没有在批注内找到是否有类似批注，若有则进行插入批注。

请严格按照以下JSON格式返回结果：
{
    "best_match_index": 数字(候选批注的编号，从1开始) 或 null,
    "similarity_score": 数字(0.0-1.0之间的相似度分数),
    "reasoning": "为什么这个批注最适合该段落？",
    "target_text_in_paragraph": "从段落中抽出的、批注应附着的原文（必须是段落的子串）",
    "has_error": true/false(是否发现明显错误),
    "error_type": "错误类型（如：日期错误、数字错误、名称错误等）",
    "context_analysis": "对同一标题下上下文的分析",
    "comment_type": "批注类型：specific_error(具体错误) 或 global_comment(全局性批注)",
    "is_title": true/false(当前段落是否为标题),
    "logic_consistent": true/false(仅对specific_error有效，原文与批注逻辑是否一致),
    "all_similarity_scores": [所有批注的相似度分数列表，按批注顺序排列],
    "modify_suggestion": "修改建议"
}

如果没有任何批注适合该段落，返回：
{
    "best_match_index": null,
    "similarity_score": 0.0,
    "reasoning": "未找到合适批注的原因",
    "target_text_in_paragraph": null,
    "has_error": false,
    "error_type": null,
    "context_analysis": "上下文分析",
    "comment_type": "none",
    "is_title": false,
    "logic_consistent": true,
    "all_similarity_scores": [所有批注的相似度分数列表，按批注顺序排列],
    "modify_suggestion": ""
}"""

    user_prompt = f"""当前段落文本：
"{target_text}"

同一标题下的上下文信息：
"{context_text}"

候选批注：
{chr(10).join(candidate_comments)}

请分析上述段落、上下文和批注，找出最匹配的组合，并按要求返回JSON。特别注意：
1. 识别明显的错误，无论文本长度如何。
2. target_text_in_paragraph 必须是当前段落中的实际文本，不能是上下文的文本。
3. 如果批注是对整个文档的宏观评价或建议，标记为global_comment类型，并且只应该匹配一次，找到最适合的段落位置插入。
4. 仿真性检查：仅对具体错误批注(specific_error)进行逻辑一致性分析。如果原文内容与批注要求逻辑一致，则不进行批注。例如：批注说"验收日期定为3日"，原文说"验收期为4月3日"，两者都是3日，逻辑一致，不批注。全局性批注不进行逻辑一致性分析。
5. 标题过滤：如果当前段落是标题（包含数字编号、章节标识等），则不进行批注匹配。
6. 批注只与正文绑定，不与标题绑定。
7. 全局性批注：找到最适合的段落位置插入，不进行逻辑一致性分析，且每个全局批注只应该匹配一次。
8. 计算所有批注的相似度分数，填入all_similarity_scores字段。"""

    data = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 800,
        "response_format": {"type": "json_object"}
    }
    
    # 增加重试机制
    max_retries = 3
    retry_delay = 2  # 秒
    
    for attempt in range(max_retries):
        try:
            # 使用更长的超时时间和重试设置
            session = requests.Session()
            session.mount('https://', requests.adapters.HTTPAdapter(
                max_retries=requests.adapters.Retry(
                    total=3,
                    backoff_factor=0.5,
                    status_forcelist=[500, 502, 503, 504]
                )
            ))
            
            response = session.post(
                DEEPSEEK_API_URL, 
                headers=headers, 
                json=data, 
                timeout=(10, 60),  # (连接超时, 读取超时)
                verify=True
            )
            
            if response.status_code != 200:
                print(f'  ❌ DeepSeek API error: {response.status_code} - {response.text}')
                if attempt < max_retries - 1:
                    print(f'  ⏳ 重试中... ({attempt + 1}/{max_retries})')
                    time.sleep(retry_delay)
                    continue
                return None
            
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            
            import json
            match_result = json.loads(content)
            
            best_index = match_result.get('best_match_index')
            similarity_score = match_result.get('similarity_score', 0.0)
            text_in_para = match_result.get('target_text_in_paragraph')
            reasoning = match_result.get('reasoning', '')
            has_error = match_result.get('has_error', False)
            error_type = match_result.get('error_type', '')
            context_analysis = match_result.get('context_analysis', '')
            comment_type = match_result.get('comment_type', 'specific_error')
            is_title = match_result.get('is_title', False)
            logic_consistent = match_result.get('logic_consistent', True)
            all_similarity_scores = match_result.get('all_similarity_scores', [])
            modify_suggestion = match_result.get('modify_suggestion', '')
            
            print(f"  AI分析结果:")
            # 显示当前批注内容（如果有匹配的话）
            if best_index is not None and 0 <= best_index - 1 < len(doc_texts):
                current_comment = doc_texts[best_index - 1]
                print(f"    当前批注内容: {current_comment}")
            else:
                print(f"    当前批注内容: 无匹配")
            print(f"    当前原文内容: {target_text}")
            print(f"    相似度: {similarity_score}")
            print(f"    所有批注相似度: {all_similarity_scores}")
            
            # 处理全局性批注和具体错误批注
            if (best_index is not None and 
                similarity_score >= threshold and 
                (has_error or comment_type == 'global_comment') and
                not is_title and
                (comment_type == 'global_comment' or not logic_consistent)):
                
                actual_index = best_index - 1
                if 0 <= actual_index < len(doc_texts):
                    matched_comment = doc_texts[actual_index]
                    
                    # 对于全局性批注，使用段落开头作为锚点
                    if comment_type == 'global_comment':
                        # 使用段落开头的实际文本，不添加省略号
                        text_in_para = target_text[:50] if len(target_text) > 50 else target_text
                        print(f"  ✓ 全局性批注，使用段落开头作为锚点: '{text_in_para}'")
                    
                    # 验证文本是否在段落中
                    if text_in_para and text_in_para in target_text:
                        return text_in_para, matched_comment, similarity_score, actual_index, reasoning, context_analysis, error_type, comment_type, is_title, logic_consistent, modify_suggestion
                    else:
                        print(f"  ❌ 模型返回的文本 '{text_in_para}' 不在段落中。")
                        return None
                else:
                    return None
            else:
                # 去掉模糊匹配，只保留精确匹配
                print(f"  ❌ 精确匹配失败，跳过此段落")
                
                if is_title:
                    print(f"  ⚠ 当前段落是标题，跳过批注")
                elif comment_type == 'specific_error' and logic_consistent:
                    print(f"  ⚠ 具体错误批注：原文与批注逻辑一致，跳过批注")
                elif not has_error and comment_type != 'global_comment':
                    print(f"  ⚠ 未发现明显错误并且非全局性批注，跳过批注")
                elif similarity_score < threshold:
                    print(f"  ⚠ 相似度 {similarity_score} 低于阈值 {threshold}")
                return None
                
        except requests.exceptions.ConnectionError as e:
            print(f'  ❌ 连接错误 (尝试 {attempt + 1}/{max_retries}): {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2  # 指数退避
                continue
            else:
                print(f'  ❌ 连接失败，已达到最大重试次数')
                return None
                
        except requests.exceptions.Timeout as e:
            print(f'  ❌ 请求超时 (尝试 {attempt + 1}/{max_retries}): {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                print(f'  ❌ 请求超时，已达到最大重试次数')
                return None
                
        except Exception as e:
            print(f'  ❌ DeepSeek API调用或解析失败 (尝试 {attempt + 1}/{max_retries}): {str(e)}')
            if attempt < max_retries - 1:
                print(f'  ⏳ 等待 {retry_delay} 秒后重试...')
                time.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                print(f'  ❌ API调用失败，已达到最大重试次数')
                return None


# 清理和标准化文本
def normalize_text(text):
    """
    清理和标准化文本，去除多余空格和换行
    """
    if not text:
        return ""
    # 去除多余空格和换行
    text = re.sub(r'\s+', ' ', text.strip())
    return text

# 确保上传目录存在
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 确保保存目录存在
SAVE_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'save-word')
if not os.path.exists(SAVE_FOLDER):
    os.makedirs(SAVE_FOLDER)
app.config['SAVE_FOLDER'] = SAVE_FOLDER

# 查找并返回本地LibreOffice的可执行文件路径，如果未安装则抛出异常。
def find_libreoffice():
    """查找 LibreOffice 可执行文件路径"""
    soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if not os.path.exists(soffice_path):
        raise Exception("未找到 LibreOffice，请确保已安装 LibreOffice")
    return soffice_path

# 使用docx2txt库将.doc文件内容读取出来，并生成新的.docx文件作为备用转换方法。
def convert_doc_to_docx_alternative(doc_path):
    """备用的.doc转换方法，尝试生成格式规范的docx"""
    try:
        # 方法1：尝试使用mammoth库（如果可用）
        try:
            import mammoth
            print("尝试使用mammoth库转换...")
            
            # mammoth主要用于docx到html的转换，但也可以尝试doc转换
            with open(doc_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                
                if html_content and len(html_content.strip()) > 0:
                    print(f"mammoth提取到HTML内容，长度: {len(html_content)}")
                    
                    # 将HTML转换为docx
                    from docx import Document
                    from bs4 import BeautifulSoup
                    
                    # 解析HTML
                    soup = BeautifulSoup(html_content, 'html.parser')
                    # 创建新的docx文档
                    doc = Document()
                    
                    
                    # 处理段落
                    for p_tag in soup.find_all(['p', 'div']):
                        if p_tag.get_text().strip():
                            p = doc.add_paragraph(p_tag.get_text().strip())
                    
                    # 如果没有找到段落，尝试处理所有文本
                    if len(doc.paragraphs) == 0:
                        text = soup.get_text()
                        if text.strip():
                            doc.add_paragraph(text.strip())
                    
                    # 保存文档
                    temp_dir = tempfile.mkdtemp()
                    converted_file = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(doc_path))[0]}.docx")
                    doc.save(converted_file)
                    
                    # 将转换后的文件保存到save-word目录
                    save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_mammoth.docx"
                    save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
                    
                    # 如果文件已存在，添加时间戳
                    if os.path.exists(save_path):
                        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_mammoth_{timestamp}.docx"
                        save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
                    
                    # 复制文件到保存目录
                    shutil.copy2(converted_file, save_path)
                    # print(f"✓ mammoth转换的docx文件已保存到: {save_path}")
                    
                    # print(f"✓ mammoth转换成功，创建了 {len(doc.paragraphs)} 个段落")
                    return converted_file
                else:
                    # print("mammoth提取的内容为空")
                    raise Exception("mammoth提取的内容为空")
                    
        except ImportError:
            print("mammoth库不可用，尝试其他方法")
        except Exception as e:
            print(f"mammoth转换失败: {e}")
        
        # 方法2：尝试使用python-docx直接读取（如果支持）
        try:
            from docx import Document
            # 注意：python-docx可能不支持直接读取.doc文件
            # 这里作为备用尝试
            doc = Document(doc_path)
            temp_dir = tempfile.mkdtemp()
            converted_file = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(doc_path))[0]}.docx")
            doc.save(converted_file)
            
            # 将转换后的文件保存到save-word目录
            save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_direct.docx"
            save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
            
            # 如果文件已存在，添加时间戳
            if os.path.exists(save_path):
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_direct_{timestamp}.docx"
                save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
            
            # 复制文件到保存目录
            shutil.copy2(converted_file, save_path)
            # print(f"✓ python-docx直接转换的docx文件已保存到: {save_path}")
            
            # print("✓ 使用python-docx直接转换成功")
            return converted_file
        except Exception as e:
            print(f"python-docx直接转换失败: {e}")
        
        # 方法3：使用docx2txt提取文本，然后创建格式规范的docx
        import docx2txt
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 读取.doc文件内容
        text = docx2txt.process(doc_path)
        print(f"从doc文件提取的文本长度: {len(text)}")
        
        # 创建新的.docx文档
        doc = Document()
        
        # 改进的文本处理：尝试保持原始格式结构
        if text.strip():
            # 按段落分割（保持原始段落结构）
            paragraphs = text.split('\n\n')  # 双换行表示段落分隔
            if len(paragraphs) == 1:
                # 如果没有双换行，尝试单换行分割
                paragraphs = text.split('\n')
            
            print(f"识别到 {len(paragraphs)} 个段落")
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    # 处理段落内的换行
                    lines = para_text.split('\n')
                    if len(lines) == 1:
                        # 单行段落
                        p = doc.add_paragraph(lines[0])
                    else:
                        # 多行段落，第一行作为段落开始
                        p = doc.add_paragraph(lines[0])
                        # 后续行添加到当前段落
                        for line in lines[1:]:
                            if line.strip():
                                p.add_run('\n' + line.strip())
        else:
            # 如果提取的文本为空，添加默认内容
            doc.add_paragraph("文档内容")
        
        # 保存到临时文件
        temp_dir = tempfile.mkdtemp()
        converted_file = os.path.join(temp_dir, f"{os.path.splitext(os.path.basename(doc_path))[0]}.docx")
        
        doc.save(converted_file)
        # print(f"备用转换完成，创建的段落数: {len(doc.paragraphs)}")
        
        # 验证转换后的文件
        try:
            with zipfile.ZipFile(converted_file, 'r') as docx:
                file_list = docx.namelist()
                if 'word/document.xml' in file_list:
                    print("✓ 备用转换生成了有效的docx文件")
                    
                    # 将转换后的文件保存到save-word目录
                    save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_backup.docx"
                    save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
                    
                    # 如果文件已存在，添加时间戳
                    if os.path.exists(save_path):
                        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                        save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_backup_{timestamp}.docx"
                        save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
                    
                    # 复制文件到保存目录
                    shutil.copy2(converted_file, save_path)
                    print(f"✓ 备用转换的docx文件已保存到: {save_path}")
                    
                    return converted_file
                else:
                    print("⚠ 备用转换生成的文件格式不正确")
                    raise Exception("备用转换生成的文件格式不正确")
        except Exception as e:
            print(f"验证转换文件失败: {e}")
            raise Exception(f"备用转换验证失败: {str(e)}")
        
    except ImportError:
        raise Exception("备用转换方法需要安装相关库: pip install docx2txt beautifulsoup4 mammoth")
    except Exception as e:
        raise Exception(f"备用转换失败: {str(e)}")

# 使用LibreOffice将.doc文件转换为.docx格式，若失败则调用备用方法进行转换。
def convert_doc_to_docx(doc_path):
    """将 .doc 文件转换为 .docx 格式"""
    temp_dir = None
    try:
        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        
        # 确保源文件存在
        if not os.path.exists(doc_path):
            raise Exception(f"源文件不存在: {doc_path}")
        
        print(f"开始转换doc文件: {doc_path}")
        print(f"临时目录: {temp_dir}")
        
        # 获取 LibreOffice 路径
        soffice_path = find_libreoffice()
        print(f"使用LibreOffice路径: {soffice_path}")
        
        # 使用更详细的LibreOffice转换命令
        cmd = [
            soffice_path,
            '--headless',
            '--convert-to', 'docx:Office Open XML Text',
            '--outdir', temp_dir,
            doc_path
        ]
        
        print(f"执行转换命令: {' '.join(cmd)}")
        
        result = subprocess.run(
            cmd, 
            capture_output=True, 
            text=True, 
            encoding='utf-8', 
            errors='ignore',
            timeout=60  # 添加超时限制
        )
        
        print(f"转换命令返回码: {result.returncode}")
        if result.stdout:
            print(f"转换输出: {result.stdout}")
        if result.stderr:
            print(f"转换错误: {result.stderr}")
        
        # 检查转换结果
        if result.returncode != 0:
            error_msg = result.stderr if result.stderr else "未知错误"
            print(f"LibreOffice转换失败，尝试备用方法: {error_msg}")
            return convert_doc_to_docx_alternative(doc_path)
            
        # 获取转换后的文件路径
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        converted_file = os.path.join(temp_dir, f"{base_name}.docx")

        # 验证转换后的文件
        if not os.path.exists(converted_file):
            print("转换后的文件不存在，检查临时目录内容:")
            # 列出临时目录中的所有文件
            for file in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, file)
                if os.path.isfile(file_path):
                    print(f"  - {file} ({os.path.getsize(file_path)} bytes)")
                else:
                    print(f"  - {file} (目录)")
            print("LibreOffice转换失败，尝试备用方法")
            return convert_doc_to_docx_alternative(doc_path)
            
        print(f"转换后的文件大小: {os.path.getsize(converted_file)} bytes")
            
        # 智能验证文件格式
        print("验证转换后的文件格式...")
        try:
            with zipfile.ZipFile(converted_file, 'r') as docx:
                file_list = docx.namelist()
                print(f"docx文件内容: {file_list}")
                
                # 检查不同的文档格式
                if 'word/document.xml' in file_list:
                    print("✓ 转换为Office Open XML格式成功")
                    
                    # 进一步验证文档内容
                    try:
                        doc_content = docx.read('word/document.xml').decode('utf-8')
                        if '<w:body>' in doc_content and '</w:body>' in doc_content:
                            print("✓ 文档结构完整")
                            
                            # 检查是否有实际内容
                            if '<w:t>' in doc_content:
                                print("✓ 文档包含文本内容")
                            else:
                                print("⚠ 文档结构完整但可能没有文本内容")
                        else:
                            print("⚠ 文档结构不完整")
                    except Exception as e:
                        print(f"验证文档内容时出错: {e}")
                        
                elif 'content.xml' in file_list:
                    print("转换为OpenDocument格式，尝试备用方法")
                    return convert_doc_to_docx_alternative(doc_path)
                elif any(f.startswith('word/') for f in file_list):
                    print("✓ 转换为部分XML格式，可能可以处理")
                else:
                    # 检查是否有其他可能的文档结构
                    print("转换后的文件格式未知，检查是否有其他结构")
                    
                    # 解压文件检查结构
                    temp_check_dir = tempfile.mkdtemp()
                    try:
                        with zipfile.ZipFile(converted_file, 'r') as check_zip:
                            check_zip.extractall(temp_check_dir)
                        
                        # 检查解压后的结构
                        print(f"解压后的结构:")
                        for root, dirs, files in os.walk(temp_check_dir):
                            level = root.replace(temp_check_dir, '').count(os.sep)
                            indent = ' ' * 2 * level
                            print(f"{indent}{os.path.basename(root)}/")
                            subindent = ' ' * 2 * (level + 1)
                            for file in files:
                                print(f"{subindent}{file}")
                        
                        # 查找可能的文档文件
                        possible_docs = []
                        for root, dirs, files in os.walk(temp_check_dir):
                            for file in files:
                                if file.endswith('.xml') and any(keyword in file.lower() for keyword in ['document', 'content', 'text']):
                                    possible_docs.append(os.path.relpath(os.path.join(root, file), temp_check_dir))
                        
                        if possible_docs:
                            print(f"找到可能的文档文件: {possible_docs}")
                            print("✓ 文件结构可能可以处理")
                        else:
                            print("没有找到任何文档文件，尝试备用方法")
                            return convert_doc_to_docx_alternative(doc_path)
                            
                    finally:
                        shutil.rmtree(temp_check_dir)
                        
        except zipfile.BadZipFile:
            print("转换后的文件不是有效的ZIP文件，尝试备用方法")
            return convert_doc_to_docx_alternative(doc_path)
                
        print(f"✓ 转换成功: {converted_file}")
        
        # 将转换后的文件保存到save-word目录
        save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}.docx"
        save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
        
        # 如果文件已存在，添加时间戳
        if os.path.exists(save_path):
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            save_filename = f"converted_{os.path.splitext(os.path.basename(doc_path))[0]}_{timestamp}.docx"
            save_path = os.path.join(app.config['SAVE_FOLDER'], save_filename)
        
        # 复制文件到保存目录
        shutil.copy2(converted_file, save_path)
        print(f"✓ 转换后的docx文件已保存到: {save_path}")
        
        return converted_file
        
    except subprocess.TimeoutExpired:
        print("LibreOffice转换超时，尝试备用方法")
        return convert_doc_to_docx_alternative(doc_path)
    except Exception as e:
        # 清理临时文件
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                print(f"已清理临时目录: {temp_dir}")
            except:
                pass
        
        # 如果LibreOffice转换失败，尝试备用方法
        print(f"LibreOffice转换失败: {str(e)}，尝试备用方法")
        return convert_doc_to_docx_alternative(doc_path)

# Flask路由，渲染首页模板。
@app.route('/')
def index():
    return render_template('index.html')

# 创建一个指定名称的Word XML元素。
def create_element(name):
    """创建 XML 元素"""
    return OxmlElement(name)

# 为指定的XML元素设置属性。
def create_attribute(element, name, value):
    """设置 XML 元素属性"""
    element.set(qn(name), value)

# 在Word文档中查找指定文本，并在该文本处插入批注引用和批注内容。
def insert_comment(doc, text, comment_text):
    """在Word文档中插入批注"""
    try:
        # 获取当前文档中已有的批注数量，用于生成新的批注ID
        comment_id = str(len(doc._comments) + 1) if hasattr(doc, '_comments') else '1'
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            if text in paragraph.text:
                # 找到包含目标文本的运行
                for run in paragraph.runs:
                    if text in run.text:
                        # 创建批注引用
                        comment_reference = create_element('w:commentReference')
                        create_attribute(comment_reference, 'w:id', comment_id)
                        
                        # 将批注引用直接插入到文本中
                        run._element.append(comment_reference)
                        
                        # 创建批注
                        comment = create_element('w:comment')
                        create_attribute(comment, 'w:id', comment_id)
                        create_attribute(comment, 'w:author', '批注系统')
                        create_attribute(comment, 'w:date', datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
                        create_attribute(comment, 'w:initials', 'PS')
                        
                        # 创建批注内容
                        p = create_element('w:p')
                        r = create_element('w:r')
                        t = create_element('w:t')
                        t.text = clean_comment_text(comment_text)
                        r.append(t)
                        p.append(r)
                        comment.append(p)
                        
                        # 初始化批注列表
                        if not hasattr(doc, '_comments'):
                            doc._comments = []
                        doc._comments.append(comment)
                        # 确保文档有comments部分
                        if not hasattr(doc._element, 'comments'):
                            comments = create_element('w:comments')
                            doc._element.append(comments)
                        else:
                            comments = doc._element.find(qn('w:comments'))
                            if comments is None:
                                comments = create_element('w:comments')
                                doc._element.append(comments)
                        # 添加批注到comments部分
                        comments.append(comment)
                        # 确保批注引用和批注内容正确关联
                        run._element.set(qn('w:commentReference'), comment_id)
                        
                        return True
        return False
    except Exception as e:
        raise Exception(f"插入批注时出错: {str(e)}")

# 从Excel文件中读取批注数据，要求Excel文件至少包含两列：要查找的文本和批注内容。
def read_comments_from_excel(excel_path):
    """从Excel文件中读取批注"""
    try:
        df = pd.read_excel(excel_path)
        if len(df.columns) < 2:
            raise Exception("Excel文件必须至少包含两列：要查找的文本和批注内容")
            
        comments = []
        # 将第一列作为要查找的文本，第二列作为批注内容
        text_col, comment_col = df.columns[0], df.columns[1]
        
        for index, row in df.iterrows():
            if pd.notna(row[text_col]) and pd.notna(row[comment_col]):
                comments.append({
                    'text': str(row[text_col]).strip(),
                    'comment': str(row[comment_col]).strip()
                })
        return comments
    except Exception as e:
        raise Exception(f"读取Excel文件时出错: {str(e)}")

# Flask路由，处理文件上传请求，将Excel中的批注插入到Word文档中，并返回处理后的文档。
@app.route('/upload_file', methods=['POST'])
def upload_file():
    if 'word_file' not in request.files or 'excel_file' not in request.files:
        return jsonify({'error': '请上传Word文档和Excel文件'}), 400
    word_file = request.files['word_file']
    excel_file = request.files['excel_file']
    if word_file.filename == '' or excel_file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    if not word_file.filename.endswith(('.doc', '.docx')) or not excel_file.filename.endswith('.xlsx'):
        return jsonify({'error': '请上传正确的文件格式'}), 400
        
    try:
        # 保存上传的文件
        word_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(word_file.filename))
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(excel_file.filename))
        
        word_file.save(word_path)
        excel_file.save(excel_path)
        
        # 记录原始文件名（用于后续查找save-word中的文件）
        original_filename = os.path.splitext(word_file.filename)[0]
        original_extension = os.path.splitext(word_file.filename)[1]
        
        # 如果是 .doc 文件，先转换为 .docx 并保存到save-word目录
        docx_path_for_processing = word_path  # 默认使用原始文件路径
        if word_path.endswith('.doc'):
            print(f"开始转换.doc文件: {word_path}")
            converted_temp_path = convert_doc_to_docx(word_path)
            
            # 验证转换后的文件
            if not os.path.exists(converted_temp_path):
                raise Exception(f"转换后的文件不存在: {converted_temp_path}")
                
            # 验证文件格式
            try:
                with zipfile.ZipFile(converted_temp_path, 'r') as docx:
                    if 'word/document.xml' not in docx.namelist():
                        raise Exception("转换后的文件格式不正确，缺少document.xml")
            except Exception as e:
                raise Exception(f"转换后的文件格式验证失败: {str(e)}")
            
            # 查找save-word目录中对应的转换后文件
            print("查找save-word目录中的转换后文件...")
            save_word_files = []
            if os.path.exists(app.config['SAVE_FOLDER']):
                for file in os.listdir(app.config['SAVE_FOLDER']):
                    if file.startswith(f"converted_{original_filename}") and file.endswith('.docx'):
                        save_word_files.append(file)
            
            if save_word_files:
                # 使用最新的文件（按修改时间排序）
                save_word_files.sort(key=lambda x: os.path.getmtime(os.path.join(app.config['SAVE_FOLDER'], x)), reverse=True)
                latest_file = save_word_files[0]
                docx_path_for_processing = os.path.join(app.config['SAVE_FOLDER'], latest_file)
                print(f"使用save-word中的文件进行批注处理: {docx_path_for_processing}")
            else:
                print("警告: 在save-word目录中未找到转换后的文件，使用临时转换文件")
                docx_path_for_processing = converted_temp_path
        else:
            # 如果是.docx文件，直接使用
            print(f"使用原始docx文件: {word_path}")
            docx_path_for_processing = word_path
            
        # 读取Excel中的批注
        comments = read_comments_from_excel(excel_path)
        print("\n=== 从Excel读取到的批注 ===")
        for i, comment in enumerate(comments):
            print(f"序号: {comment['text']}")
            print(f"批注: {comment['comment']}")
            print("---")
            
        # 保存修改后的文档
        original_name = os.path.splitext(os.path.basename(docx_path_for_processing))[0]
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'output_{original_name}.docx')
        
        # 检查输出文件是否已存在，如果存在则删除
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception as e:
                print(f"警告: 无法删除已存在的输出文件: {str(e)}")
                # 如果无法删除，使用新的文件名
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'output_{uuid.uuid4()}.docx')
        
        # 确保输出目录存在且有写入权限
        try:
            if not os.path.exists(app.config['UPLOAD_FOLDER']):
                os.makedirs(app.config['UPLOAD_FOLDER'])
            # 测试写入权限
            test_file = os.path.join(app.config['UPLOAD_FOLDER'], 'test_write.tmp')
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
        except Exception as e:
            print(f"错误: 无法写入输出目录: {str(e)}")
            return jsonify({'error': '无法写入输出目录，请检查权限设置'}), 500
        
        try:
            # 第一步：复制用于处理的Word文件
            print(f"\n=== 第一步：复制Word文件用于批注处理 ===")
            print(f"源文件: {docx_path_for_processing}")
            copied_word_path = os.path.join(app.config['UPLOAD_FOLDER'], f'copied_{uuid.uuid4()}.docx')
            shutil.copy2(docx_path_for_processing, copied_word_path)
            print(f"Word文件已复制到: {copied_word_path}")

            # 第二步：使用新的集成处理函数（包含文档预处理、题号分类和上下文感知批注匹配）
            print("\n=== 第二步：执行集成批注处理 ===")
            add_comments_to_docx_xml(copied_word_path, comments, output_path)
            print(f"\n文档已保存到: {output_path}")
            
            # 直接使用docx文件作为最终输出，不再转换为doc
            final_output_path = output_path
            print(f"✓ 最终输出文件: {final_output_path}")
            
            # 自动打开处理完成的文件
            try:
                print(f"\n=== 自动打开处理完成的文件 ===")
                # 使用系统默认程序打开文件
                if os.name == 'nt':  # Windows
                    os.startfile(final_output_path)
                elif os.name == 'posix':  # macOS 和 Linux
                    subprocess.run(['open', final_output_path], check=True)
                print(f"✓ 已自动打开文件: {final_output_path}")
            except Exception as e:
                print(f"⚠️ 自动打开文件失败: {str(e)}")
                print(f"请手动打开文件: {final_output_path}")
            
            # 清理临时文件
            try:
                if os.path.exists(copied_word_path) and copied_word_path != output_path:
                    os.remove(copied_word_path)
            except Exception as e:
                print(f"警告: 清理复制的Word文件失败: {str(e)}")
            try:
                if os.path.exists(word_path) and word_path != output_path:
                    os.remove(word_path)
            except Exception as e:
                print(f"警告: 清理原始Word文件失败: {str(e)}")
            try:
                if os.path.exists(excel_path):
                    os.remove(excel_path)
            except Exception as e:
                print(f"警告: 清理Excel文件失败: {str(e)}")
            
            # 清理临时和原始文件，只保留最终输出文件
            for f in [word_path, excel_path, copied_word_path]:
                if os.path.exists(f) and f != final_output_path:
                    try:
                        os.remove(f)
                    except Exception as e:
                        print(f"警告: 删除文件 {f} 失败: {e}")
            # 清理 save-word 目录下的所有文件
            if os.path.exists(app.config['SAVE_FOLDER']):
                for f in os.listdir(app.config['SAVE_FOLDER']):
                    file_path = os.path.join(app.config['SAVE_FOLDER'], f)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    except Exception as e:
                        print(f"警告: 删除中间文件 {file_path} 失败: {e}")
            
            return jsonify({
                'message': f'文件处理成功',
                'output_file': os.path.basename(final_output_path)
            })
            
        except Exception as e:
            print(f"\n=== 错误信息 ===")
            print(str(e))
            # 如果输出文件已存在，尝试删除
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            return jsonify({'error': str(e)}), 500
        
    except Exception as e:
        print(f"\n=== 错误信息 ===")
        print(str(e))
        return jsonify({'error': str(e)}), 500

# 清理批注文本中的特殊字符，确保XML格式正确，避免解析错误。
def clean_comment_text(text):
    """清理批注文本，确保特殊字符被正确转义"""
    return (text.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;"))
            #    .replace('"', "&quot;")
            #    .replace("'", "&apos;"))

def parse_docx_blocks(paragraphs):
    """
    解析多级标题，构建区块树结构，返回每个内容段落所属的最小区块。
    区块结构：{'title': 标题文本, 'level': 级别, 'start': 段落索引, 'end': 段落索引, 'children': [], 'content_indices': []}
    返回：block_tree, para_to_block_idx 映射
    """
    # 更宽松的标题正则，支持数字、中文数字、括号、字母等
    title_pattern = re.compile(
        r'^((\d+(?:\.\d+)*|[一二三四五六七八九十百千万]+|[(（][一二三四五六七八九十百千万\d]+[)）]|[A-Za-z])([\.、．\s)])+)'
    )
    block_stack = []
    block_tree = []
    para_to_block_idx = {}
    default_block = {'title': '默认区块', 'level': 0, 'start': 0, 'end': 0, 'children': [], 'content_indices': []}
    in_default = True

    for idx, para in enumerate(paragraphs):
        text = ''.join(t.text for t in para.findall('.//w:t', namespaces={'w': W_NS}) if t.text).strip()
        m = title_pattern.match(text)
        if m:
            # 识别标题级别
            level = m.group(1).count('.') + 1
            block = {'title': text, 'level': level, 'start': idx, 'end': idx, 'children': [], 'content_indices': []}
            # 结束默认区块
            if in_default and default_block['content_indices']:
                default_block['end'] = idx - 1
                block_tree.append(default_block)
                in_default = False
            # 找到父区块
            while block_stack and block_stack[-1]['level'] >= level:
                finished = block_stack.pop()
                finished['end'] = idx - 1
            if block_stack:
                block_stack[-1]['children'].append(block)
            else:
                block_tree.append(block)
            block_stack.append(block)
        else:
            if in_default:
                default_block['content_indices'].append(idx)
                para_to_block_idx[idx] = default_block
            elif block_stack:
                block_stack[-1]['content_indices'].append(idx)
                para_to_block_idx[idx] = block_stack[-1]
    last_idx = len(paragraphs) - 1
    if in_default and default_block['content_indices']:
        default_block['end'] = last_idx
        block_tree.append(default_block)
    while block_stack:
        finished = block_stack.pop()
        finished['end'] = last_idx
    return block_tree, para_to_block_idx

def add_comments_to_docx_xml(docx_path, comments, output_path):
    print(f"\n=== 开始处理docx文件（多级区块+内容批注） ===")
    print(f"输入文件: {docx_path}")
    print(f"输出文件: {output_path}")
    
    if not os.path.exists(docx_path):
        raise Exception(f"输入文件不存在: {docx_path}")
    
    temp_dir = 'temp_docx_' + str(uuid.uuid4())
    os.makedirs(temp_dir, exist_ok=True)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
    except Exception as e:
        raise Exception(f"文件解压失败: {str(e)}")

    doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
    comments_xml_path = os.path.join(temp_dir, 'word', 'comments.xml')
    w_ns = W_NS
    parser = etree.XMLParser(remove_blank_text=True, recover=True)
    doc_tree = etree.parse(doc_xml_path, parser)
    doc_root = doc_tree.getroot()
    paragraphs = doc_root.findall('.//w:p', namespaces={'w': w_ns})

    # 解析区块树和内容段落归属
    block_tree, para_to_block_idx = parse_docx_blocks(paragraphs)

    # 创建或加载comments.xml
    if os.path.exists(comments_xml_path):
        comments_tree = etree.parse(comments_xml_path, parser)
        comments_root = comments_tree.getroot()
    else:
        MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
        W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
        WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
        nsmap = {'w': w_ns, 'mc': MC_NS, 'w14': W14_NS, 'w15': W15_NS, 'wp14': WP14_NS}
        comments_root = etree.Element(f'{{{w_ns}}}comments', nsmap=nsmap)
        comments_root.set(f"{{{MC_NS}}}Ignorable", "w14 w15 wp14")
        comments_tree = etree.ElementTree(comments_root)

    current_id = 0
    comment_texts = [item['comment'].strip() for item in comments if item.get('comment', '').strip()]
    
    # 跟踪已匹配的批注索引
    matched_comment_indices = set()
    
    # 跟踪已插入的全局批注，避免重复插入
    inserted_global_comments = set()
    
    # 发处理段落批注匹配
    print(f"\n=== 开始并发批注处理 ===")
    start_time = time.time()
    
    # 动态调整并发参数
    total_paragraphs = len([p for p in paragraphs if p.findall('.//w:t', namespaces={'w': w_ns})])
    total_comments = len(comment_texts)
    adjust_concurrent_config(total_paragraphs, total_comments)
    
    # 检查是否启用并发处理
    # 所有文档都使用并发处理以提高速度，但增强错误处理
    if CONCURRENT_CONFIG['enabled']:
        print(f"✓ 启用并发处理模式")
        print(f"  批次大小: {CONCURRENT_CONFIG['batch_size']}")
        print(f"  最大并发数: {CONCURRENT_CONFIG['max_workers']}")
        print(f"  相似度阈值: {CONCURRENT_CONFIG['threshold']}")
        
        # 创建段落批次
        paragraph_batches = create_paragraph_batches(paragraphs, para_to_block_idx, comment_texts, batch_size=CONCURRENT_CONFIG['batch_size'])
        
        # 并发处理所有批次
        all_results = []
        for batch_idx, batch in enumerate(paragraph_batches):
            print(f"\n--- 处理批次 {batch_idx + 1}/{len(paragraph_batches)} ---")
            batch_results = process_paragraph_batch(batch, comment_texts, threshold=CONCURRENT_CONFIG['threshold'], max_workers=CONCURRENT_CONFIG['max_workers'])
            all_results.extend(batch_results)
    else:
        print(f"⚠ 使用串行处理模式")
        # 串行处理（保持原有逻辑作为备用）
        all_results = []
        for para_idx, para in enumerate(paragraphs):
            para_text = ''.join(t.text for t in para.findall('.//w:t', namespaces={'w': w_ns}) if t.text).strip()
            if para_idx not in para_to_block_idx or not para_text or not comment_texts:
                continue
            block = para_to_block_idx[para_idx]
            content_indices = block['content_indices']
            idx_in_block = content_indices.index(para_idx)
            context_text = ""
            if idx_in_block > 0:
                prev_idx = content_indices[idx_in_block - 1]
                prev_text = ''.join(t.text for t in paragraphs[prev_idx].findall('.//w:t', namespaces={'w': w_ns}) if t.text).strip()
                if prev_text:
                    context_text += f"同区块前文：{prev_text}\n"
            if idx_in_block < len(content_indices) - 1:
                next_idx = content_indices[idx_in_block + 1]
                next_text = ''.join(t.text for t in paragraphs[next_idx].findall('.//w:t', namespaces={'w': w_ns}) if t.text).strip()
                if next_text:
                    context_text += f"同区块后文：{next_text}"
            
            match = deepseek_context_aware_match(para_text, comment_texts, context_text, threshold=CONCURRENT_CONFIG['threshold'])
            if match:
                all_results.append({
                    'para_idx': para_idx,
                    'para': para,
                    'para_text': para_text,
                    'match': match,
                    'success': True
                })
            else:
                all_results.append({
                    'para_idx': para_idx,
                    'para': para,
                    'para_text': para_text,
                    'match': None,
                    'success': False
                })
    
    # 处理并发结果
    for result in all_results:
        if not result['success']:
            continue
            
        para_idx = result['para_idx']
        para = result['para']
        para_text = result['para_text']
        text_to_find, original_comment, _, comment_idx, reasoning, context_analysis, error_type, comment_type, is_title, logic_consistent, modify_suggestion = result['match']
        comment_content = original_comment
        comment_id = str(current_id)
        
        # 记录已匹配的批注索引
        matched_comment_indices.add(comment_idx)
        
        # 对于全局性批注，检查是否已经插入过
        if comment_type == 'global_comment':
            if comment_idx in inserted_global_comments:
                print(f"  ⚠ 全局批注 {comment_idx} 已插入过，跳过重复插入")
                continue
            inserted_global_comments.add(comment_idx)
        
        print(f"  ✓ 段落 {para_idx} 匹配到批注: '{original_comment}'")
        print(f"    应附着于文本: '{text_to_find}'")
        print(f"    匹配原因: '{reasoning}'")
        print(f"    上下文分析: '{context_analysis}'")
        print(f"    错误类型: '{error_type}'")
        print(f"    批注类型: '{comment_type}'")
        print(f"    是否为标题: {is_title}")
        if comment_type == 'specific_error':
            print(f"    原文与批注逻辑是否一致: {logic_consistent}")
        else:
            print(f"    原文与批注逻辑是否一致: 全局性批注不进行逻辑分析")
        if comment_type == 'global_comment':
            print(f"    这是全局性批注，将在批注内容前添加【全局】标识")

        if is_title:
            print(f"    ⚠ 当前段落是标题，将跳过批注")
        if comment_type == 'specific_error' and logic_consistent:
            print(f"    ⚠ 具体错误批注：原文与批注逻辑一致，将跳过批注")
        
        # === 修正批注锚点插入逻辑，完全模拟Word ===
        # 1. 合并需要批注的文本区间到一个 run
        # 2. 在 run 前插入 commentRangeStart，在 run 后插入 commentRangeEnd 和 commentReference
        runs = para.findall('.//w:r', namespaces={'w': w_ns})
        run_map = []
        full_para_text = ""
        for i, run in enumerate(runs):
            run_text = "".join(t.text for t in run.findall('.//w:t', namespaces={'w': w_ns}) if t.text)
            if run_text:
                run_map.append({'run': run, 'text': run_text, 'start': len(full_para_text)})
                full_para_text += run_text
        start_pos = full_para_text.find(text_to_find)
        if start_pos == -1:
            print(f"    ⚠ 在段落 {para_idx} 中未精确定位到模型返回的文本，跳过。")
            continue
        end_pos = start_pos + len(text_to_find)
        # 找到 run 区间
        start_run_idx, end_run_idx = -1, -1
        for i, r_info in enumerate(run_map):
            run_end_pos = r_info['start'] + len(r_info['text'])
            if start_run_idx == -1 and start_pos < run_end_pos:
                start_run_idx = i
            if end_run_idx == -1 and end_pos <= run_end_pos:
                end_run_idx = i
                break
        if start_run_idx == -1: continue
        if end_run_idx == -1: end_run_idx = len(run_map) - 1
        # 合并区间内所有 run 的文本到第一个 run
        merged_text = ""
        for i in range(start_run_idx, end_run_idx + 1):
            merged_text += run_map[i]['text']
        # 清空区间内所有 run 的 <w:t>
        for i in range(start_run_idx, end_run_idx + 1):
            for t in run_map[i]['run'].findall('.//w:t', namespaces={'w': w_ns}):
                t.text = ""
        # 在第一个 run 的第一个 <w:t> 填入合并后的文本
        first_run = run_map[start_run_idx]['run']
        t_list = first_run.findall('.//w:t', namespaces={'w': w_ns})
        if t_list:
            t_list[0].text = merged_text
        # 插入 commentRangeStart
        parent = first_run.getparent()
        comment_start_node = etree.Element(W('commentRangeStart'))
        comment_start_node.set(W('id'), comment_id)
        
        # 使用线程锁保护XML插入操作
        with xml_lock:
            try:
                # 重新获取当前索引，因为可能在并发环境下发生变化
                current_first_run_index = parent.index(first_run) if first_run in parent else -1
                if current_first_run_index >= 0:
                    parent.insert(current_first_run_index, comment_start_node)
                else:
                    # 如果找不到索引，添加到开头
                    parent.insert(0, comment_start_node)
            except Exception as e:
                print(f"    ⚠ 插入commentRangeStart失败: {str(e)}")
                # 尝试添加到段落开头
                try:
                    parent.insert(0, comment_start_node)
                except Exception as e2:
                    print(f"    ❌ 插入commentRangeStart到开头也失败: {str(e2)}")
                    # 最后尝试添加到末尾
                    try:
                        parent.append(comment_start_node)
                    except Exception as e3:
                        print(f"    ❌ 插入commentRangeStart到末尾也失败: {str(e3)}")
                        print(f"    ❌ 跳过段落 {para_idx} 的批注插入")
                        continue
        
        # 插入 commentRangeEnd
        last_run = run_map[end_run_idx]['run']
        comment_end_node = etree.Element(W('commentRangeEnd'))
        comment_end_node.set(W('id'), comment_id)
        
        with xml_lock:
            try:
                # 重新获取当前索引
                current_last_run_index = parent.index(last_run) if last_run in parent else -1
                if current_last_run_index >= 0:
                    parent.insert(current_last_run_index + 1, comment_end_node)
                else:
                    # 如果找不到索引，添加到末尾
                    parent.append(comment_end_node)
            except Exception as e:
                print(f"    ⚠ 插入commentRangeEnd失败: {str(e)}")
                # 尝试添加到段落末尾
                try:
                    parent.append(comment_end_node)
                except Exception as e2:
                    print(f"    ❌ 插入commentRangeEnd到末尾也失败: {str(e2)}")
                    # 如果还是失败，跳过这个批注
                    print(f"    ❌ 跳过段落 {para_idx} 的批注插入")
                    continue
        
        # 插入 commentReference
        comment_ref_run = etree.Element(W('r'))
        comment_ref_node = etree.Element(W('commentReference'))
        comment_ref_node.set(W('id'), comment_id)
        comment_ref_run.append(comment_ref_node)
        
        with xml_lock:
            try:
                # 重新获取当前索引
                current_last_run_index = parent.index(last_run) if last_run in parent else -1
                if current_last_run_index >= 0:
                    parent.insert(current_last_run_index + 2, comment_ref_run)
                else:
                    # 如果找不到索引，添加到末尾
                    parent.append(comment_ref_run)
            except Exception as e:
                # 如果插入失败，尝试添加到段落末尾
                print(f"    ⚠ 插入commentReference失败，尝试添加到段落末尾: {str(e)}")
                try:
                    parent.append(comment_ref_run)
                except Exception as e2:
                    print(f"    ❌ 插入commentReference到末尾也失败: {str(e2)}")
                    # 如果还是失败，跳过这个批注
                    print(f"    ❌ 跳过段落 {para_idx} 的批注插入")
                    continue
        # 创建批注内容到comments.xml
        comment_elem = etree.Element(W('comment'))
        comment_elem.set(W('id'), comment_id)
        comment_elem.set(W('author'), '批注系统')
        comment_elem.set(W('date'), datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
        # 第一行：批注内容
        p1 = etree.Element(W('p'))
        p1.set(f'{{{W14_NS}}}paraId', gen_para_id())
        pPr = etree.Element(W('pPr'))
        pStyle = etree.Element(W('pStyle'))
        pStyle.set(W('val'), 'CommentText')
        pPr.append(pStyle)
        p1.append(pPr)
        r_ann = etree.Element(W('r'))
        annotationRef = etree.Element(W('annotationRef'))
        r_ann.append(annotationRef)
        p1.append(r_ann)
        r1 = etree.Element(W('r'))
        t1 = etree.Element(W('t'))
        # 只显示原始批注内容，不包含匹配原因
        if comment_type == 'global_comment':
            comment_content = "【全局】" + comment_content
        t1.text = clean_comment_text(comment_content)
        r1.append(t1)
        p1.append(r1)
        comment_elem.append(p1)
        
        # 第二行：空行（回车）
        p_empty = etree.Element(W('p'))
        p_empty.set(f'{{{W14_NS}}}paraId', gen_para_id())
        r_empty = etree.Element(W('r'))
        t_empty = etree.Element(W('t'))
        t_empty.text = ""  # 空文本，形成空行
        r_empty.append(t_empty)
        p_empty.append(r_empty)
        comment_elem.append(p_empty)
        
        # 第三行：匹配原因
        if reasoning:
            p3 = etree.Element(W('p'))
            p3.set(f'{{{W14_NS}}}paraId', gen_para_id())
            r3 = etree.Element(W('r'))
            t3 = etree.Element(W('t'))
            t3.text = clean_comment_text("批注原因：" + reasoning)
            r3.append(t3)
            p3.append(r3)
            comment_elem.append(p3)

        # 第四行：空行
        p_empty = etree.Element(W('p'))
        p_empty.set(f'{{{W14_NS}}}paraId', gen_para_id())
        r_empty = etree.Element(W('r'))
        t_empty = etree.Element(W('t'))
        t_empty.text = ""  # 空文本，形成空行
        r_empty.append(t_empty)
        p_empty.append(r_empty)
        comment_elem.append(p_empty)

        # 第五行：修改建议
        p4 = etree.Element(W('p'))
        p4.set(f'{{{W14_NS}}}paraId', gen_para_id())
        r4 = etree.Element(W('r'))
        t4 = etree.Element(W('t'))
        t4.text = clean_comment_text("修改建议：" + (modify_suggestion or ""))
        r4.append(t4)
        p4.append(r4)
        comment_elem.append(p4)

        # 使用线程锁保护comments.xml的写入
        with xml_lock:
            comments_root.append(comment_elem)
        print(f"    ✓ 成功插入批注。")
        current_id += 1
    
    end_time = time.time()
    processing_time = end_time - start_time
    print(f"\n✓ 并发批注处理完成，总耗时: {processing_time:.2f} 秒")
    print(f"✓ 成功处理 {len([r for r in all_results if r['success']])} 个段落")
    
    # 性能统计
    successful_matches = len([r for r in all_results if r['success']])
    total_processed = len(all_results)
    success_rate = (successful_matches / total_processed * 100) if total_processed > 0 else 0
    avg_time_per_para = processing_time / total_processed if total_processed > 0 else 0
    
    print(f"\n=== 性能统计 ===")
    print(f"  总段落数: {total_processed}")
    print(f"  成功匹配: {successful_matches}")
    print(f"  成功率: {success_rate:.1f}%")
    print(f"  总耗时: {processing_time:.2f} 秒")
    print(f"  平均每段耗时: {avg_time_per_para:.3f} 秒")
    
    if CONCURRENT_CONFIG['enabled']:
        estimated_serial_time = total_processed * avg_time_per_para
        speedup = estimated_serial_time / processing_time if processing_time > 0 else 1
        # print(f"  预估串行耗时: {estimated_serial_time:.2f} 秒")
        # print(f"  加速比: {speedup:.2f}x")
        
        # 性能分析和建议
        print(f"\n=== 性能分析 ===")
        if speedup < 1.5:
            print(f"  ⚠️ 加速比较低 ({speedup:.2f}x)，可能的原因：")
            print(f"    - API响应时间较长 ({avg_time_per_para:.1f}秒/段)")
            print(f"    - 并发数可能不够 (当前: {CONCURRENT_CONFIG['max_workers']})")
            print(f"    - 网络延迟较高")
            print(f"  建议：")
            print(f"    - 增加并发数到 8-10")
            print(f"    - 检查网络连接")
            print(f"    - 考虑降低API超时时间")
        elif speedup < 3.0:
            print(f"  ✓ 加速比中等 ({speedup:.2f}x)，有优化空间")
            print(f"  建议：")
            print(f"    - 可以尝试增加并发数")
            print(f"    - 优化批次大小")
        else:
            print(f"  🎉 加速比优秀 ({speedup:.2f}x)，并发效果很好！")
        
        # 计算API调用效率
        total_api_calls = total_processed
        api_calls_per_second = total_api_calls / processing_time if processing_time > 0 else 0
        print(f"  API调用效率: {api_calls_per_second:.2f} 次/秒")
        
        if api_calls_per_second < 0.5:
            print(f"  ⚠️ API调用效率较低，建议增加并发数")
        elif api_calls_per_second < 1.0:
            print(f"  ✓ API调用效率中等")
        else:
            print(f"  🎉 API调用效率很高！")
 
    # 处理未匹配的批注：合并为一个批注（每条回车分隔）
    # print(f"\n=== 处理未匹配的批注 ===")
    # unmatched_comments = []
    # for i, comment in enumerate(comments):
    #     if i not in matched_comment_indices and comment.get('comment', '').strip():
    #         unmatched_comments.append((i, comment))
    
    # if unmatched_comments:
    #     print(f"找到 {len(unmatched_comments)} 个未匹配的批注，合并为一个批注插入")
    #     # 找到文档的body元素
    #     body = doc_root.find('.//w:body', namespaces={'w': w_ns})
    #     if body is None:
    #         print("警告：未找到文档body元素")
    #     else:
    #         # 插入一个可见文本作为所有未匹配批注的锚点
    #         anchor_para = etree.Element(W('p'))
    #         comment_id = str(current_id)
    #         comment_start_node = etree.Element(W('commentRangeStart'))
    #         comment_start_node.set(W('id'), comment_id)
    #         anchor_para.append(comment_start_node)
    #         run = etree.Element(W('r'))
    #         t = etree.Element(W('t'))
    #         t.text = "【未匹配成功批注】"
    #         run.append(t)
    #         anchor_para.append(run)
    #         comment_end_node = etree.Element(W('commentRangeEnd'))
    #         comment_end_node.set(W('id'), comment_id)
    #         anchor_para.append(comment_end_node)
    #         comment_ref_run = etree.Element(W('r'))
    #         comment_ref_node = etree.Element(W('commentReference'))
    #         comment_ref_node.set(W('id'), comment_id)
    #         comment_ref_run.append(comment_ref_node)
    #         anchor_para.append(comment_ref_run)
    #         body.append(anchor_para)
    #         # 合并所有未匹配内容为一个批注
    #         comment_elem = etree.Element(W('comment'))
    #         comment_elem.set(W('id'), comment_id)
    #         comment_elem.set(W('author'), '批注系统')
    #         comment_elem.set(W('date'), datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
    #         for comment_idx, comment_data in unmatched_comments:
    #             comment_content = comment_data['comment'].strip()
    #             comment_text = comment_data.get('text', '').strip()
    #             # 每条为一个段落
    #             p = etree.Element(W('p'))
    #             r = etree.Element(W('r'))
    #             t = etree.Element(W('t'))
    #             t.text = clean_comment_text(f"【未匹配成功】{comment_content}")
    #             r.append(t)
    #             p.append(r)
    #             comment_elem.append(p)
    #             # 原始序号
    #             if comment_text:
    #                 p2 = etree.Element(W('p'))
    #                 r2 = etree.Element(W('r'))
    #                 t2 = etree.Element(W('t'))
    #                 t2.text = clean_comment_text("原始序号：" + comment_text)
    #                 r2.append(t2)
    #                 p2.append(r2)
    #                 comment_elem.append(p2)
    #             # 状态说明
    #             p3 = etree.Element(W('p'))
    #             r3 = etree.Element(W('r'))
    #             t3 = etree.Element(W('t'))
    #             t3.text = clean_comment_text("状态：未找到匹配的文档内容，作为独立批注添加")
    #             r3.append(t3)
    #             p3.append(r3)
    #             comment_elem.append(p3)
    #         comments_root.append(comment_elem)
    #         current_id += 1
    else:
        print("✓ 所有批注都已成功匹配")
    
    # 保存XML
    doc_tree.write(doc_xml_path, xml_declaration=True, encoding='utf-8', standalone='yes')
    comments_tree.write(comments_xml_path, xml_declaration=True, encoding='utf-8', standalone='yes')
    
    # 第三步：打包为新的docx文件
    print("\n=== 第三步：打包为新的docx文件 ===")
    
    # 确保 word/_rels/document.xml.rels 文件存在并包含批注关系
    rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
    rels_dir = os.path.dirname(rels_path)
    os.makedirs(rels_dir, exist_ok=True)
    
    # 初始化 create_new_rels 变量
    create_new_rels = False
    
    if os.path.exists(rels_path):
        try:
            rels_tree = etree.parse(rels_path)
            rels_root = rels_tree.getroot()
            # 检查是否已有 comments 关系
            has_comments_rel = any(
                rel.get('Type') == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
                for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            )
            if not has_comments_rel:
                # 添加 comments 关系
                rel = etree.Element('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
                rel.set('Id', 'rIdComments')
                rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
                rel.set('Target', 'comments.xml')
                rels_root.append(rel)
                rels_tree.write(rels_path, xml_declaration=True, encoding='utf-8', standalone='yes')
        except Exception as e:
            print(f"警告: 修改现有rels文件失败: {str(e)}")
            create_new_rels = True
    else:
        create_new_rels = True
    
    if create_new_rels:
        # 创建新的rels文件
        rels_root = etree.Element('Relationships', nsmap={None: 'http://schemas.openxmlformats.org/package/2006/relationships'})
        
        # 添加基本关系
        rel = etree.Element('Relationship')
        rel.set('Id', 'rId1')
        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles')
        rel.set('Target', 'styles.xml')
        rels_root.append(rel)
        
        # 添加批注关系
        rel = etree.Element('Relationship')
        rel.set('Id', 'rIdComments')
        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
        rel.set('Target', 'comments.xml')
        rels_root.append(rel)
        
        rels_tree = etree.ElementTree(rels_root)
        rels_tree.write(rels_path, xml_declaration=True, encoding='utf-8', standalone='yes')
    
    # === 打包前，确保 [Content_Types].xml 包含 comments.xml 类型声明 ===
    content_types_path = os.path.join(temp_dir, '[Content_Types].xml')
    if os.path.exists(content_types_path):
        try:
            parser = etree.XMLParser(remove_blank_text=True, recover=True)
            tree = etree.parse(content_types_path, parser)
            root = tree.getroot()
            ns = {'ct': 'http://schemas.openxmlformats.org/package/2006/content-types'}
            # 检查是否已有 comments.xml 类型声明
            has_comments_override = any(
                el.get('PartName') == '/word/comments.xml' and
                el.get('ContentType') == 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
                for el in root.findall('ct:Override', namespaces=ns)
            )
            if not has_comments_override:
                # 插入 Override 节点
                override = etree.Element('{http://schemas.openxmlformats.org/package/2006/content-types}Override')
                override.set('PartName', '/word/comments.xml')
                override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
                # 插入到最后一个 Override 之后
                last_override = None
                for el in root.findall('ct:Override', namespaces=ns):
                    last_override = el
                if last_override is not None:
                    last_override.addnext(override)
                else:
                    root.append(override)
                tree.write(content_types_path, xml_declaration=True, encoding='utf-8', standalone='yes')
                print('✓ 已自动补全 [Content_Types].xml 中的 comments.xml 类型声明')
        except Exception as e:
            print(f'警告: 自动补全 [Content_Types].xml 失败: {e}')
    
    # 删除旧的输出文件
    if os.path.exists(output_path):
        os.remove(output_path)
    
    # 打包为docx文件
    try:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 打包前，删除 temp_dir 根目录下的 document.xml 和 comments.xml（只保留 word/ 目录下的）
        for fname in ['document.xml', 'comments.xml']:
            fpath = os.path.join(temp_dir, fname)
            if os.path.exists(fpath):
                os.remove(fpath)
        
        file_order = [
            '[Content_Types].xml',
            '_rels/.rels',
            'word/_rels/document.xml.rels',
            'word/document.xml',
            'word/comments.xml'
        ]
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
            written_files = set()
            packed_files = []  # 用于统计打包的文件
            for file_name in file_order:
                file_path = os.path.join(temp_dir, file_name)
                arcname = file_name.replace('\\', '/')  # 统一为正斜杠
                if os.path.exists(file_path):
                    docx.write(file_path, arcname)
                    written_files.add(arcname)
                    packed_files.append(arcname)
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    file_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path, temp_dir).replace('\\', '/')  # 统一为正斜杠
                    if arcname in written_files:
                        continue
                    if filename in ['document.xml', 'comments.xml'] and not arcname.startswith('word/'):
                        continue
                    docx.write(file_path, arcname)
                    written_files.add(arcname)
                    packed_files.append(arcname)
            # 打包结束后输出统计
            print("=== 实际打包进docx的文件列表 ===")
            for f in packed_files:
                print(f)
            print("=== 文件数量统计 ===")
            for fname, count in Counter([os.path.basename(f) for f in packed_files]).items():
                print(f"{fname}: {count} 个")
            print(f"总文件数: {len(packed_files)}")
        
        # 验证输出文件
        try:
            with zipfile.ZipFile(output_path, 'r') as test_zip:
                if 'word/document.xml' not in test_zip.namelist():
                    raise Exception("输出文件缺少document.xml")
                if 'word/comments.xml' not in test_zip.namelist():
                    raise Exception("输出文件缺少comments.xml")
                if 'word/_rels/document.xml.rels' not in test_zip.namelist():
                    raise Exception("输出文件缺少document.xml.rels")
            print("✓ docx文件打包成功并验证通过")
        except Exception as e:
            raise Exception(f"输出文件验证失败: {str(e)}")
            
    except Exception as e:
        raise Exception(f"打包docx文件失败: {str(e)}")

    # 清理临时文件夹
    try:
        shutil.rmtree(temp_dir)
    except Exception as e:
        print(f"警告: 清理临时文件夹时出错: {str(e)}")
    
    print(f"✓ 最终输出文件: {output_path}")
    
    return output_path

# 使用LibreOffice将.docx文件转换为.doc格式，并重命名输出文件到指定路径。
def convert_docx_to_doc(docx_path, doc_path):
    """将.docx文件转换为.doc格式（使用LibreOffice）"""
    try:
        print(f"开始将docx转换为doc: {docx_path} -> {doc_path}")
        
        # 验证输入文件
        if not os.path.exists(docx_path):
            raise Exception(f"输入文件不存在: {docx_path}")
        
        print(f"输入文件大小: {os.path.getsize(docx_path)} bytes")
        
        # 获取LibreOffice路径
        soffice_path = find_libreoffice()
        print(f"使用LibreOffice: {soffice_path}")
        
        # 准备转换命令
        output_dir = os.path.dirname(doc_path)
        cmd = [
            soffice_path,
            '--headless',
            '--convert-to', 'doc',
            '--outdir', output_dir,
            docx_path
        ]
        
        print(f"执行命令: {' '.join(cmd)}")
        
        # 执行转换
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        print(f"转换命令返回码: {result.returncode}")
        if result.stdout:
            print(f"转换输出: {result.stdout}")
        if result.stderr:
            print(f"转换错误: {result.stderr}")
        
        if result.returncode != 0:
            raise Exception(f'LibreOffice转换docx到doc失败: {result.stderr}')
        
        # LibreOffice会自动命名输出文件，需重命名
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        generated_doc = os.path.join(output_dir, base_name + '.doc')
        
        print(f"LibreOffice生成的文件: {generated_doc}")
        print(f"目标文件: {doc_path}")
        
        # 检查生成的文件是否存在
        if not os.path.exists(generated_doc):
            # 列出输出目录中的所有文件
            print(f"输出目录 {output_dir} 中的文件:")
            for file in os.listdir(output_dir):
                file_path = os.path.join(output_dir, file)
                if os.path.isfile(file_path):
                    print(f"  - {file} ({os.path.getsize(file_path)} bytes)")
                else:
                    print(f"  - {file} (目录)")
            raise Exception(f"LibreOffice未生成预期的doc文件: {generated_doc}")
        
        # 重命名文件
        if generated_doc != doc_path:
            print(f"重命名文件: {generated_doc} -> {doc_path}")
            os.rename(generated_doc, doc_path)
        
        # 验证最终文件
        if os.path.exists(doc_path):
            print(f"✓ 转换成功: {doc_path}")
            print(f"最终文件大小: {os.path.getsize(doc_path)} bytes")
        else:
            raise Exception(f"最终文件不存在: {doc_path}")
            
    except subprocess.TimeoutExpired:
        raise Exception("LibreOffice转换超时")
    except Exception as e:
        raise Exception(f"转换docx到doc失败: {str(e)}")

def gen_para_id():
    # 生成8位十六进制字符串
    return f"{random.randint(0, 0xFFFFFFFF):08X}"

# 主程序入口，启动Flask应用，默认在本地开发环境运行。
if __name__ == '__main__':
    app.run(debug=True)